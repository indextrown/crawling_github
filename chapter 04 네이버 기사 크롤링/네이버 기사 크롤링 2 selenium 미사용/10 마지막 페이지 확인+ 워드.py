import requests
from bs4 import BeautifulSoup
import time
import pyautogui
from docx import Document

# 사용자 입력
keyword = pyautogui.prompt("검색어를 입력하시오:")
# pyautogui.prompt는 리턴 데이터 기본 값은 문자열임 그래서 int로 강제형변환
lastpage = int(pyautogui.prompt("몇 페이지까지 크롤링 할까요?"))

# 1. 워드 문서 생성하기
document = Document()

page_num = 1
# 여러 페이지 가져오기
for i in range(1, lastpage * 10, 10):
    print(f"{page_num}페이지 크롤링 중입니다.===========")
    response = requests.get(
        f"https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}&start={i}")
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')

    # 뉴스 기사 div 10개 추출
    articles = soup.select("div.news_info")

    #
    for article in articles:
        links = article.select("a.info")

        # 링크가 2개 이상이면
        if len(links) >= 2:
            # 두번째 링크의 href를 추출한다
            url = links[1].attrs['href']

            # headers={'user-agent':'Mozila/5.0'} 봇으로 접근 가능하게 함
            response = requests.get(url, headers={'user-agent': 'Mozila/5.0'})
            html = response.text
            soup_sub = BeautifulSoup(html, 'html.parser')

            content = soup_sub.select_one("#dic_area")
            title = soup_sub.select_one(".media_end_head_title")
            # 만약 연애 뉴스 라면 //response => 링크가 리다이렉션인지(사이트 주소가 바뀌는지) 확인하기 위해 response.url 로 만듬

            # 만약 연애 뉴스 라면 //response => 링크가 리다이렉션인지(사이트 주소가 바뀌는지) 확인하기 위해 response.url 로 만듬
            if "entertain" in response.url:
                # 연예뉴스 사이트 css선택자 사용
                title = soup_sub.select_one(".end_tit")
                content = soup_sub.select_one("#articeBody")
            # 그게 아니고 스포츠 뉴스 url 이면
            elif "sports" in response.url:
                # 스포츠뉴스 사이트 css선택자 사용
                title = soup_sub.select_one("h4.title")
                content = soup_sub.select_one("#newsEndContents")
                # 본문 내용안에 불필요한 div, p 삭제
                divs = content.select("div")
                for div in divs:
                    div.decompose()
                paragraphs = content.select("p")
                for p in paragraphs:
                    p.decompose()


            # 그게 아니면
            else:
                # 일반뉴스 사이트 css선택자 사용
                title = soup_sub.select_one(".media_end_head_title")
                content = soup_sub.select_one("#dic_area")
                brs = content.select("br")
                for br in brs:
                    br.decompose()
                brs = content.select("br")
                for br in brs:
                    br.decompose()
# strip 공백제거
            print("=======링크======= \n", url)
            # 공백 제거 .text.strip()
            print("=======제목======= \n", title.text.strip())

            print("=======본문======= \n", content.text.strip())

            # 2. 워드에 제목, 링크, 본문 데이터 추가하기
            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())
            # time.sleep(0.3)
    #
    # 마지막 페이지 여부 확인하기
    lastpage = soup.select_one(".btn_next").attrs['aria-disabled']
    if lastpage == 'true':
        print("마지막 페이지 입니다.")
        # 가장 가까운 반복문을 빠져나간다
        break
    page_num = page_num + 1

# 3. 워드 문서 저장하기
document.save(f"{keyword}.docx")
